"""Manuscript import API endpoints for importing existing prose into the project."""

import io
import re
import logging
import asyncio
from datetime import datetime
from fastapi import APIRouter, HTTPException, UploadFile, File, Form, BackgroundTasks
from pydantic import BaseModel, Field
from typing import List, Optional

from app.config import settings
from app.utils.file_utils import write_json_file, read_json_file
from app.utils.backup import create_snapshot
from app.services.memory_service import memory_service
from app.services.entity_service import get_entity_service

router = APIRouter()
logger = logging.getLogger(__name__)


class ManuscriptPreview(BaseModel):
    """Preview of imported manuscript."""
    filename: str
    format: str
    total_words: int
    total_chars: int
    text_preview: str  # First 2000 chars
    full_text: str
    detected_chapters: Optional[List["ChapterSplit"]] = None  # Pre-detected from heading styles


class ChapterSplit(BaseModel):
    """A chapter split from the manuscript."""
    chapter_number: int
    title: str
    content: str
    word_count: int


class ManuscriptSplitPreview(BaseModel):
    """Preview of manuscript split into chapters."""
    chapters: List[ChapterSplit]
    total_chapters: int
    total_words: int


class ImportAsSceneRequest(BaseModel):
    """Request to import manuscript text as a scene in edit mode."""
    text: str = Field(..., description="The manuscript text to import")
    scene_title: str = Field(..., description="Title for the new scene")
    chapter_id: str = Field(..., description="Chapter to create the scene in")
    scene_number: Optional[int] = None


class ImportResult(BaseModel):
    """Result of manuscript import."""
    scene_id: str
    title: str
    word_count: int
    edit_mode: bool
    message: str


class BulkImportRequest(BaseModel):
    """Request to import multiple chapters as scenes."""
    chapters: List[ChapterSplit]
    chapter_id: str | None = None  # Chapter to create scenes in (legacy)
    act_id: str | None = None  # Act to create chapters in (optional)
    enable_edit_mode: bool = True  # Enable edit mode for each scene
    create_chapters: bool = True  # Create chapters from manuscript structure


class BulkImportResult(BaseModel):
    """Result of bulk manuscript import."""
    chapters_created: int = 0
    scenes_created: int = 0
    total_words: int = 0
    chapter_ids: List[str] = []
    scene_ids: List[str] = []
    message: str
    deep_import_started: bool = False


class DeepImportStatus(BaseModel):
    """Status of a deep import extraction job."""
    project_id: str
    series_id: Optional[str]
    total_scenes: int
    scenes_extracted: int
    current_scene: Optional[str]
    status: str  # "running", "completed", "failed"
    error: Optional[str] = None
    started_at: str
    completed_at: Optional[str] = None


# In-memory tracking of deep import jobs (would use Redis/DB in production)
_deep_import_jobs: dict[str, DeepImportStatus] = {}


async def run_deep_extraction(
    project_id: str,
    series_id: str,
    scene_ids: List[str],
    model: Optional[str] = None
):
    """
    Background task to run extraction on all imported scenes.

    Performs two types of extraction:
    1. Entity extraction (characters, world) - once per book, saved to series level
    2. Memory extraction (plot events, state changes) - per scene, saved to memory layer

    Runs sequentially to avoid rate limits.
    """
    job_id = f"{project_id}:{series_id}"

    _deep_import_jobs[job_id] = DeepImportStatus(
        project_id=project_id,
        series_id=series_id,
        total_scenes=len(scene_ids),
        scenes_extracted=0,
        current_scene=None,
        status="running",
        started_at=datetime.utcnow().isoformat()
    )

    scenes_dir = settings.scenes_dir(project_id)
    project_data = await read_json_file(settings.project_dir(project_id) / "project.json")
    book_number = project_data.get("book_number", 1)

    try:
        # Phase 1: Gather all prose and run entity extraction (once per book)
        logger.info(f"Deep import: gathering prose from {len(scene_ids)} scenes...")
        all_prose = []
        for scene_id in scene_ids:
            scene_file = scenes_dir / f"{scene_id}.json"
            if scene_file.exists():
                scene_data = await read_json_file(scene_file)
                prose = scene_data.get("original_prose") or scene_data.get("prose")
                if prose:
                    all_prose.append(prose)

        if all_prose:
            combined_prose = "\n\n---\n\n".join(all_prose)
            logger.info(f"Deep import: extracting entities from {len(combined_prose)} chars...")
            _deep_import_jobs[job_id].current_scene = "extracting_entities"

            try:
                entity_service = get_entity_service()
                entity_result = await entity_service.extract_and_save_entities(
                    series_id=series_id,
                    book_id=project_id,
                    book_number=book_number,
                    prose=combined_prose,
                    model=model
                )
                char_stats = entity_result.get("characters", {})
                world_stats = entity_result.get("world_elements", {})
                logger.info(
                    f"Deep import: entities extracted - "
                    f"characters: {char_stats.get('created', 0)} created, {char_stats.get('updated', 0)} updated | "
                    f"world: {world_stats.get('created', 0)} created, {world_stats.get('updated', 0)} updated"
                )
            except Exception as e:
                logger.warning(f"Entity extraction failed: {e}")
                # Continue with memory extraction even if entity extraction fails

            await asyncio.sleep(2)  # Pause before memory extraction

        # Phase 2: Per-scene memory extraction
        logger.info(f"Deep import: starting per-scene memory extraction...")
        for i, scene_id in enumerate(scene_ids):
            _deep_import_jobs[job_id].current_scene = scene_id

            # Load scene data
            scene_file = scenes_dir / f"{scene_id}.json"
            if not scene_file.exists():
                continue

            scene_data = await read_json_file(scene_file)
            prose = scene_data.get("original_prose") or scene_data.get("prose")

            if not prose:
                continue

            # Get chapter info for context
            chapter_id = scene_data.get("chapter_id")
            chapter_title = None
            chapter_number = None
            if chapter_id:
                chapter_file = settings.project_dir(project_id) / "chapters" / f"{chapter_id}.json"
                if chapter_file.exists():
                    chapter_data = await read_json_file(chapter_file)
                    chapter_title = chapter_data.get("title")
                    chapter_number = chapter_data.get("chapter_number")

            # Run memory extraction
            try:
                await memory_service.extract_from_scene(
                    series_id=series_id,
                    book_id=project_id,
                    scene_id=scene_id,
                    prose=prose,
                    scene_title=scene_data.get("title"),
                    chapter_title=chapter_title,
                    book_number=book_number,
                    chapter_number=chapter_number,
                    scene_number=scene_data.get("scene_number"),
                    model=model
                )

                _deep_import_jobs[job_id].scenes_extracted = i + 1

                # Small delay between extractions to be nice to API
                await asyncio.sleep(1)

            except Exception as e:
                logger.warning(f"Memory extraction failed for scene {scene_id}: {e}")
                # Continue with other scenes even if one fails

        # Phase 3: Generate memory summaries
        try:
            logger.info(f"Deep import: generating memory summaries...")
            await memory_service.generate_summaries(series_id, model=model)
        except Exception as e:
            logger.warning(f"Summary generation failed: {e}")

        _deep_import_jobs[job_id].status = "completed"
        _deep_import_jobs[job_id].completed_at = datetime.utcnow().isoformat()
        _deep_import_jobs[job_id].current_scene = None
        logger.info(f"Deep import completed for {project_id}")

    except Exception as e:
        _deep_import_jobs[job_id].status = "failed"
        _deep_import_jobs[job_id].error = str(e)
        _deep_import_jobs[job_id].completed_at = datetime.utcnow().isoformat()
        logger.error(f"Deep import failed for {project_id}: {e}")


def ensure_project_exists(project_id: str):
    """Check that project exists, raise 404 if not."""
    if not settings.project_dir(project_id).exists():
        raise HTTPException(status_code=404, detail=f"Project not found: {project_id}")


def validate_chapter_exists(project_id: str, chapter_id: str):
    """Validate that a chapter exists in the project."""
    chapters_dir = settings.project_dir(project_id) / "chapters"
    chapter_file = chapters_dir / f"{chapter_id}.json"
    if not chapter_file.exists():
        raise HTTPException(status_code=400, detail=f"Chapter not found: {chapter_id}")


def slugify(text: str) -> str:
    """Convert text to a URL-friendly slug."""
    text = text.lower().strip()
    text = re.sub(r'[^\w\s-]', '', text)
    text = re.sub(r'[-\s]+', '-', text)
    return text[:50]  # Limit length


def convert_docx_to_text(file_bytes: bytes) -> str:
    """Convert .docx file to plain text using mammoth."""
    try:
        import mammoth
        result = mammoth.extract_raw_text(io.BytesIO(file_bytes))
        return result.value
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="mammoth library not installed. Run: pip install mammoth"
        )
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to parse .docx file: {str(e)}")


def detect_chapters_from_docx(file_bytes: bytes) -> List[ChapterSplit] | None:
    """
    Try to detect chapters from .docx heading styles using python-docx.

    If the docx has Heading 1 and/or Heading 2 styles, uses them to split:
    - Heading 1 = Act/Part (used as grouping context, not a chapter itself)
    - Heading 2 = Chapter
    - If only Heading 1 is present with no Heading 2, treats H1 as chapters.

    Returns None if python-docx isn't installed or no heading styles found,
    so the caller can fall back to regex-based detection.
    """
    try:
        import docx as docxlib
    except ImportError:
        return None

    try:
        doc = docxlib.Document(io.BytesIO(file_bytes))
    except Exception:
        return None

    # Scan for heading styles
    h1_count = 0
    h2_count = 0
    for p in doc.paragraphs:
        if p.style.name == 'Heading 1' and p.text.strip():
            h1_count += 1
        elif p.style.name == 'Heading 2' and p.text.strip():
            h2_count += 1

    if h1_count == 0 and h2_count == 0:
        return None  # No heading styles, fall back to regex

    # Determine which heading level marks chapters
    if h2_count > 0:
        # Two-level structure: H1 = parts/acts, H2 = chapters
        chapter_heading = 'Heading 2'
    else:
        # Single-level: H1 = chapters
        chapter_heading = 'Heading 1'

    # Collect chapters by splitting on the chapter heading level
    chapters = []
    current_title = None
    current_lines = []
    chapter_num = 0

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue

        if p.style.name == chapter_heading:
            # Flush previous chapter
            if current_title is not None and current_lines:
                content = '\n\n'.join(current_lines)
                word_count = len(content.split())
                if word_count > 0:
                    chapters.append(ChapterSplit(
                        chapter_number=chapter_num,
                        title=current_title,
                        content=content,
                        word_count=word_count,
                    ))
            chapter_num += 1
            current_title = text
            current_lines = []
        elif p.style.name == 'Heading 1' and chapter_heading == 'Heading 2':
            # Part/Act heading — skip as structural marker (not a chapter)
            # but include as context if it has content before first chapter
            continue
        else:
            # Body text — accumulate into current chapter
            if current_title is not None:
                current_lines.append(text)

    # Flush last chapter
    if current_title is not None and current_lines:
        content = '\n\n'.join(current_lines)
        word_count = len(content.split())
        if word_count > 0:
            chapters.append(ChapterSplit(
                chapter_number=chapter_num,
                title=current_title,
                content=content,
                word_count=word_count,
            ))

    if len(chapters) >= 2:
        logger.info(f"Detected {len(chapters)} chapters from docx heading styles ({chapter_heading})")
        return chapters

    return None  # Not enough chapters found, fall back to regex


WORD_TO_NUM = {
    'one': 1, 'two': 2, 'three': 3, 'four': 4, 'five': 5,
    'six': 6, 'seven': 7, 'eight': 8, 'nine': 9, 'ten': 10,
    'eleven': 11, 'twelve': 12, 'thirteen': 13, 'fourteen': 14,
    'fifteen': 15, 'sixteen': 16, 'seventeen': 17, 'eighteen': 18,
    'nineteen': 19, 'twenty': 20, 'twenty-one': 21, 'twenty-two': 22,
    'twenty-three': 23, 'twenty-four': 24, 'twenty-five': 25,
    'twenty-six': 26, 'twenty-seven': 27, 'twenty-eight': 28,
    'twenty-nine': 29, 'thirty': 30
}

ROMAN_TO_NUM = {
    'i': 1, 'ii': 2, 'iii': 3, 'iv': 4, 'v': 5,
    'vi': 6, 'vii': 7, 'viii': 8, 'ix': 9, 'x': 10,
    'xi': 11, 'xii': 12, 'xiii': 13, 'xiv': 14, 'xv': 15,
    'xvi': 16, 'xvii': 17, 'xviii': 18, 'xix': 19, 'xx': 20
}


def parse_chapter_number(text: str, fallback: int) -> int:
    """Convert various chapter number formats to int."""
    text = text.strip().lower()

    # Try as digit
    try:
        return int(text)
    except ValueError:
        pass

    # Try as word
    if text in WORD_TO_NUM:
        return WORD_TO_NUM[text]

    # Try as roman numeral
    if text in ROMAN_TO_NUM:
        return ROMAN_TO_NUM[text]

    return fallback


def detect_chapter_splits(text: str) -> List[ChapterSplit]:
    """
    Attempt to split text into chapters based on common patterns.

    Tries patterns in order of specificity:
    1. "Chapter X" / "CHAPTER X" / "Chapter One" / "Chapter 01" (with optional title)
    2. "Part X" / "PART X"
    3. Standalone number words on own line: "One", "Two", "Three"
    4. Standalone digits on own line: "1", "2", "3"
    5. Roman numerals on own line: "I", "II", "III"
    """

    # Normalize line endings to \n
    text = text.replace('\r\n', '\n').replace('\r', '\n')

    # Patterns to try, in order of specificity
    patterns = [
        # "Chapter 1" / "Chapter 01" / "Chapter One" / "CHAPTER 1: Title"
        # Allows trailing whitespace, optional title after : or -
        (r'^(?:#\s*)?(?:CHAPTER|Chapter)\s+(\d+|[A-Za-z]+(?:-[A-Za-z]+)?)(?:\s*[-:]\s*(.+?))?\s*$', 'chapter'),
        # "Part 1" / "Part One"
        (r'^(?:PART|Part)\s+(\d+|[A-Za-z]+(?:-[A-Za-z]+)?)(?:\s*[-:]\s*(.+?))?\s*$', 'part'),
        # Standalone number word on its own line (preceded by blank line)
        (r'(?:\n\s*\n\s*\n+)(One|Two|Three|Four|Five|Six|Seven|Eight|Nine|Ten|Eleven|Twelve|Thirteen|Fourteen|Fifteen|Sixteen|Seventeen|Eighteen|Nineteen|Twenty(?:-(?:One|Two|Three|Four|Five|Six|Seven|Eight|Nine))?)(?:\s*\n)', 'word'),
        # Standalone digit(s) on own line (preceded by blank lines)
        (r'(?:\n\s*\n\s*\n+)(\d{1,2})(?:\s*\n)', 'digit'),
        # Roman numerals on own line (preceded by blank lines)
        (r'(?:\n\s*\n\s*\n+)(I{1,3}|IV|VI{0,3}|IX|XI{0,3}|XIV|XV|XVI{0,3}|XIX|XX)(?:\s*\n)', 'roman'),
    ]

    matches = []
    pattern_type = None

    for pattern, ptype in patterns:
        compiled = re.compile(pattern, re.MULTILINE | re.IGNORECASE)
        found = list(compiled.finditer(text))
        if len(found) >= 2:  # Need at least 2 chapters to consider it valid
            matches = found
            pattern_type = ptype
            break

    if not matches:
        # No chapter markers found, return entire text as one chapter
        word_count = len(text.split())
        return [ChapterSplit(
            chapter_number=1,
            title="Imported Manuscript",
            content=text.strip(),
            word_count=word_count
        )]

    chapters = []
    for i, match in enumerate(matches):
        chapter_num_str = match.group(1)
        chapter_num = parse_chapter_number(chapter_num_str, i + 1)

        # Get title if present (only some patterns have group 2)
        try:
            title = match.group(2).strip() if match.group(2) else f"Chapter {chapter_num}"
        except IndexError:
            title = f"Chapter {chapter_num}"

        # Get content (from end of this match to start of next, or end of text)
        start = match.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        content = text[start:end].strip()

        word_count = len(content.split())

        chapters.append(ChapterSplit(
            chapter_number=chapter_num,
            title=title,
            content=content,
            word_count=word_count
        ))

    # Handle content before first chapter (prologue/intro)
    first_match_start = matches[0].start()
    if first_match_start > 100:  # More than 100 chars before first chapter
        prologue_content = text[:first_match_start].strip()
        prologue_words = len(prologue_content.split())
        if prologue_words > 50:  # Only if substantial
            chapters.insert(0, ChapterSplit(
                chapter_number=0,
                title="Prologue",
                content=prologue_content,
                word_count=prologue_words
            ))

    return chapters


@router.post("/upload", response_model=ManuscriptPreview)
async def upload_manuscript(
    project_id: str,
    file: UploadFile = File(...)
):
    """
    Upload a manuscript file and get a preview of its contents.

    Supports: .docx, .txt, .md files

    Returns the full text content for client-side processing.
    """
    ensure_project_exists(project_id)

    # Check file extension
    filename = file.filename or "unknown"
    ext = filename.lower().split('.')[-1] if '.' in filename else ''

    if ext not in ['docx', 'txt', 'md', 'markdown']:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file format: .{ext}. Supported: .docx, .txt, .md"
        )

    try:
        # Read file content
        file_bytes = await file.read()

        # For docx, try heading-style chapter detection first
        detected_chapters = None
        if ext == 'docx':
            detected_chapters = detect_chapters_from_docx(file_bytes)
            text = convert_docx_to_text(file_bytes)
            file_format = "docx"
        else:
            # Plain text or markdown
            text = file_bytes.decode('utf-8')
            file_format = "txt" if ext == 'txt' else "markdown"

        # Calculate stats
        total_words = len(text.split())
        total_chars = len(text)
        preview = text[:2000] + ("..." if len(text) > 2000 else "")

        return ManuscriptPreview(
            filename=filename,
            format=file_format,
            total_words=total_words,
            total_chars=total_chars,
            text_preview=preview,
            full_text=text,
            detected_chapters=detected_chapters,
        )

    except HTTPException:
        raise
    except UnicodeDecodeError:
        raise HTTPException(status_code=400, detail="Failed to decode file. Ensure it's UTF-8 encoded.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to process file: {str(e)}")


@router.post("/split", response_model=ManuscriptSplitPreview)
async def split_manuscript(
    project_id: str,
    text: str = Form(...)
):
    """
    Split manuscript text into chapters based on chapter markers.

    Looks for patterns like "Chapter 1", "CHAPTER ONE", "Chapter 1: Title", etc.

    Returns a list of chapter splits that can be imported as scenes.
    """
    ensure_project_exists(project_id)

    if not text or len(text.strip()) < 10:
        raise HTTPException(status_code=400, detail="Text is too short to process")

    try:
        chapters = detect_chapter_splits(text)
        total_words = sum(c.word_count for c in chapters)

        return ManuscriptSplitPreview(
            chapters=chapters,
            total_chapters=len(chapters),
            total_words=total_words
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to split manuscript: {str(e)}")


@router.post("/import-scene", response_model=ImportResult)
async def import_as_scene(
    project_id: str,
    request: ImportAsSceneRequest
):
    """
    Import manuscript text as a single scene in edit mode.

    Creates a new scene with the imported text set as original_prose,
    ready for the critique-revision loop.
    """
    ensure_project_exists(project_id)
    validate_chapter_exists(project_id, request.chapter_id)

    try:
        scenes_dir = settings.scenes_dir(project_id)
        scenes_dir.mkdir(parents=True, exist_ok=True)

        # Generate scene ID
        scene_id = slugify(request.scene_title)
        if not scene_id:
            scene_id = f"scene-{datetime.utcnow().timestamp():.0f}"

        filepath = scenes_dir / f"{scene_id}.json"

        # Ensure unique ID
        if filepath.exists():
            scene_id = f"{scene_id}-{int(datetime.utcnow().timestamp())}"
            filepath = scenes_dir / f"{scene_id}.json"

        # Auto-assign scene_number if not provided
        scene_number = request.scene_number
        if scene_number is None:
            existing_count = 0
            for f in scenes_dir.glob("*.json"):
                try:
                    data = await read_json_file(f)
                    if data.get("chapter_id") == request.chapter_id:
                        existing_count += 1
                except:
                    continue
            scene_number = existing_count + 1

        now = datetime.utcnow()
        word_count = len(request.text.split())

        # Create scene in edit mode
        scene_data = {
            "id": scene_id,
            "title": request.scene_title,
            "outline": f"[Imported manuscript - {word_count} words]",
            "chapter_id": request.chapter_id,
            "scene_number": scene_number,
            "character_ids": [],
            "world_context_ids": [],
            "previous_scene_ids": [],
            "tags": ["imported"],
            "additional_notes": None,
            "tone": None,
            "pov": None,
            "target_length": None,
            "is_canon": False,
            "prose": None,
            "summary": None,
            "edit_mode": True,
            "original_prose": request.text,
            "edit_mode_started_at": now.isoformat(),
            "created_at": now.isoformat(),
            "updated_at": now.isoformat()
        }

        await write_json_file(filepath, scene_data)

        return ImportResult(
            scene_id=scene_id,
            title=request.scene_title,
            word_count=word_count,
            edit_mode=True,
            message=f"Scene created with {word_count} words in edit mode. Ready for critique."
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to create scene: {str(e)}")


@router.post("/import-bulk", response_model=BulkImportResult)
async def import_bulk_scenes(
    project_id: str,
    request: BulkImportRequest
):
    """
    Import multiple chapter splits, creating chapters and scenes.

    For each detected chapter in the manuscript:
    - Creates a Chapter in the project structure
    - Creates a Scene within that chapter with the prose in edit mode
    """
    ensure_project_exists(project_id)

    if not request.chapters:
        raise HTTPException(status_code=400, detail="No chapters provided")

    # Create snapshot before bulk import
    await create_snapshot(project_id, f"pre-import-{len(request.chapters)}-chapters", reason="pre-import")

    try:
        chapters_dir = settings.project_dir(project_id) / "chapters"
        chapters_dir.mkdir(parents=True, exist_ok=True)
        scenes_dir = settings.scenes_dir(project_id)
        scenes_dir.mkdir(parents=True, exist_ok=True)

        # Get starting chapter number
        existing_chapters = list(chapters_dir.glob("*.json"))
        starting_chapter_num = len(existing_chapters) + 1

        chapter_ids = []
        scene_ids = []
        total_words = 0
        now = datetime.utcnow()

        for i, manuscript_chapter in enumerate(request.chapters):
            chapter_number = starting_chapter_num + i
            word_count = len(manuscript_chapter.content.split())
            total_words += word_count

            # Create Chapter
            chapter_id = slugify(manuscript_chapter.title)
            if not chapter_id:
                chapter_id = f"chapter-{chapter_number}"

            chapter_filepath = chapters_dir / f"{chapter_id}.json"

            # Ensure unique chapter ID
            if chapter_filepath.exists():
                chapter_id = f"{chapter_id}-{int(datetime.utcnow().timestamp())}"
                chapter_filepath = chapters_dir / f"{chapter_id}.json"

            chapter_data = {
                "id": chapter_id,
                "title": manuscript_chapter.title,
                "chapter_number": chapter_number,
                "act_id": request.act_id,  # May be None
                "description": f"Imported from manuscript ({word_count} words)",
                "notes": None,
                "created_at": now.isoformat(),
                "updated_at": now.isoformat()
            }

            await write_json_file(chapter_filepath, chapter_data)
            chapter_ids.append(chapter_id)

            # Create Scene within the chapter
            scene_id = f"{chapter_id}-scene-1"
            scene_filepath = scenes_dir / f"{scene_id}.json"

            # Ensure unique scene ID
            if scene_filepath.exists():
                scene_id = f"{scene_id}-{int(datetime.utcnow().timestamp())}"
                scene_filepath = scenes_dir / f"{scene_id}.json"

            scene_data = {
                "id": scene_id,
                "title": f"{manuscript_chapter.title}",
                "outline": f"[Imported manuscript chapter - {word_count} words]",
                "chapter_id": chapter_id,
                "scene_number": 1,
                "character_ids": [],
                "world_context_ids": [],
                "previous_scene_ids": [],
                "tags": ["imported"],
                "additional_notes": None,
                "tone": None,
                "pov": None,
                "target_length": None,
                "is_canon": False,
                "prose": None,
                "summary": None,
                "created_at": now.isoformat(),
                "updated_at": now.isoformat()
            }

            # Add edit mode fields if requested
            if request.enable_edit_mode:
                scene_data["edit_mode"] = True
                scene_data["original_prose"] = manuscript_chapter.content
                scene_data["edit_mode_started_at"] = now.isoformat()

            await write_json_file(scene_filepath, scene_data)
            scene_ids.append(scene_id)

        return BulkImportResult(
            chapters_created=len(chapter_ids),
            scenes_created=len(scene_ids),
            total_words=total_words,
            chapter_ids=chapter_ids,
            scene_ids=scene_ids,
            message=f"Created {len(chapter_ids)} chapters and {len(scene_ids)} scenes with {total_words} total words."
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to import manuscript: {str(e)}")


@router.get("/deep-import-status")
async def get_deep_import_status(project_id: str):
    """
    Get the status of a deep import extraction job.

    Returns extraction progress including scenes processed and current status.
    The project_id comes from the route prefix.
    """
    # Find job for this project
    for job_id, status in _deep_import_jobs.items():
        if status.project_id == project_id:
            return status

    raise HTTPException(status_code=404, detail="No deep import job found for this project")


@router.post("/import-full", response_model=BulkImportResult)
async def import_full_manuscript(
    project_id: str,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    enable_edit_mode: bool = Form(True),
    deep_import: bool = Form(False)
):
    """
    One-shot manuscript import: upload file → detect chapters → create all.

    This is the most efficient endpoint for automated imports.
    Accepts .docx, .txt, .md files. Auto-detects chapter markers
    (Chapter 1, CHAPTER ONE, etc.) and creates chapters + scenes.

    If deep_import=True, also runs memory extraction on each scene (requires
    project to be in a series). Extraction runs in background - check status
    via GET /deep-import-status/{project_id}.

    Returns a small summary - no large text in response.
    """
    ensure_project_exists(project_id)

    # Check file extension
    filename = file.filename or "unknown"
    ext = filename.lower().split('.')[-1] if '.' in filename else ''

    if ext not in ['docx', 'txt', 'md', 'markdown']:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file format: .{ext}. Supported: .docx, .txt, .md"
        )

    try:
        # Read and convert file
        file_bytes = await file.read()

        # For docx files, try heading-style detection first
        chapters = None
        if ext == 'docx':
            chapters = detect_chapters_from_docx(file_bytes)

        if chapters is None:
            # Fall back to plain text regex detection
            if ext == 'docx':
                text = convert_docx_to_text(file_bytes)
            else:
                text = file_bytes.decode('utf-8')
            chapters = detect_chapter_splits(text)

        if not chapters:
            raise HTTPException(status_code=400, detail="No content found in file")

        # Create snapshot before import
        await create_snapshot(project_id, f"pre-import-{filename}", reason="pre-import")

        # Set up directories
        chapters_dir = settings.project_dir(project_id) / "chapters"
        chapters_dir.mkdir(parents=True, exist_ok=True)
        scenes_dir = settings.scenes_dir(project_id)
        scenes_dir.mkdir(parents=True, exist_ok=True)

        # Get starting chapter number
        existing_chapters = list(chapters_dir.glob("*.json"))
        starting_chapter_num = len(existing_chapters) + 1

        chapter_ids = []
        scene_ids = []
        total_words = 0
        now = datetime.utcnow()

        for i, manuscript_chapter in enumerate(chapters):
            chapter_number = starting_chapter_num + i
            word_count = manuscript_chapter.word_count
            total_words += word_count

            # Create Chapter
            chapter_id = slugify(manuscript_chapter.title)
            if not chapter_id:
                chapter_id = f"chapter-{chapter_number}"

            chapter_filepath = chapters_dir / f"{chapter_id}.json"

            if chapter_filepath.exists():
                chapter_id = f"{chapter_id}-{int(now.timestamp())}"
                chapter_filepath = chapters_dir / f"{chapter_id}.json"

            chapter_data = {
                "id": chapter_id,
                "title": manuscript_chapter.title,
                "chapter_number": chapter_number,
                "act_id": None,
                "description": f"Imported from {filename} ({word_count} words)",
                "notes": None,
                "created_at": now.isoformat(),
                "updated_at": now.isoformat()
            }

            await write_json_file(chapter_filepath, chapter_data)
            chapter_ids.append(chapter_id)

            # Create Scene
            scene_id = f"{chapter_id}-scene-1"
            scene_filepath = scenes_dir / f"{scene_id}.json"

            if scene_filepath.exists():
                scene_id = f"{scene_id}-{int(now.timestamp())}"
                scene_filepath = scenes_dir / f"{scene_id}.json"

            scene_data = {
                "id": scene_id,
                "title": manuscript_chapter.title,
                "outline": f"[Imported from {filename} - {word_count} words]",
                "chapter_id": chapter_id,
                "scene_number": 1,
                "character_ids": [],
                "world_context_ids": [],
                "previous_scene_ids": [],
                "tags": ["imported"],
                "additional_notes": None,
                "tone": None,
                "pov": None,
                "target_length": None,
                "is_canon": False,
                "prose": None,
                "summary": None,
                "created_at": now.isoformat(),
                "updated_at": now.isoformat()
            }

            if enable_edit_mode:
                scene_data["edit_mode"] = True
                scene_data["original_prose"] = manuscript_chapter.content
                scene_data["edit_mode_started_at"] = now.isoformat()

            await write_json_file(scene_filepath, scene_data)
            scene_ids.append(scene_id)

        # Handle deep import if requested
        deep_import_started = False
        if deep_import:
            # Get project to check for series
            project_file = settings.project_dir(project_id) / "project.json"
            project_data = await read_json_file(project_file)
            series_id = project_data.get("series_id")

            if series_id:
                # Queue deep extraction as background task
                background_tasks.add_task(
                    run_deep_extraction,
                    project_id=project_id,
                    series_id=series_id,
                    scene_ids=scene_ids
                )
                deep_import_started = True
                logger.info(f"Deep import started for {project_id} in series {series_id}")
            else:
                logger.warning(f"Deep import requested but project {project_id} is not in a series")

        message = f"Imported '{filename}': {len(chapter_ids)} chapters, {total_words} words"
        if deep_import_started:
            message += f". Deep extraction started for {len(scene_ids)} scenes (check /manuscript/deep-import-status)"

        return BulkImportResult(
            chapters_created=len(chapter_ids),
            scenes_created=len(scene_ids),
            total_words=total_words,
            chapter_ids=chapter_ids,
            scene_ids=scene_ids,
            message=message,
            deep_import_started=deep_import_started
        )

    except HTTPException:
        raise
    except UnicodeDecodeError:
        raise HTTPException(status_code=400, detail="Failed to decode file. Ensure UTF-8 encoding.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Import failed: {str(e)}")
